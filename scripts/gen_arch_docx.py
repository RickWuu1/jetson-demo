#!/usr/bin/env python3
"""Generate jetson-demo architecture diagram and append to Astar 项目.docx."""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

REPO = Path(__file__).resolve().parent.parent
OUT_DIR = REPO / "outputs" / "doc_assets"
DOCX_SRC = Path(r"c:\Users\dawn\Desktop\Astar项目.docx")
DOCX_DST = DOCX_SRC
BACKUP = DOCX_SRC.with_suffix(".docx.bak")


def _box(ax, xy, w, h, text, fc, ec="#333333", fontsize=9, weight="normal"):
    x, y = xy
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2,
        y + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        weight=weight,
        color="#111111",
        wrap=True,
    )
    return (x + w / 2, y, x + w / 2, y + h)


def _edge(box: tuple[float, float, float, float], side: str) -> tuple[float, float]:
    """box = (x, y, w, h). Return anchor on given side."""
    x, y, w, h = box
    if side == "top":
        return (x + w / 2, y + h)
    if side == "bottom":
        return (x + w / 2, y)
    if side == "left":
        return (x, y + h / 2)
    if side == "right":
        return (x + w, y + h / 2)
    raise ValueError(side)


def _draw_polyline(
    ax,
    points: list[tuple[float, float]],
    *,
    color: str = "#2C2C2C",
    lw: float = 2.2,
    style: str = "-",
    arrow_end: bool = True,
    zorder: int = 2,
) -> None:
    """Draw a complete orthogonal polyline; optional arrowhead on the last segment."""
    if len(points) < 2:
        return
    xs = [p[0] for p in points]
    ys = [p[1] for p in points]
    ax.plot(
        xs,
        ys,
        color=color,
        linewidth=lw,
        linestyle=style,
        solid_capstyle="round",
        solid_joinstyle="round",
        zorder=zorder,
    )
    if not arrow_end:
        return
    ax.annotate(
        "",
        xy=points[-1],
        xytext=points[-2],
        arrowprops=dict(
            arrowstyle="-|>",
            color=color,
            lw=lw,
            linestyle=style,
            shrinkA=0,
            shrinkB=0,
            mutation_scale=16,
        ),
        zorder=zorder + 1,
    )


def _link_v(
    ax,
    x: float,
    y_from: float,
    y_to: float,
    *,
    color: str = "#2C2C2C",
    lw: float = 2.2,
    style: str = "-",
) -> None:
    """Vertical link (y_from > y_to means downward)."""
    _draw_polyline(ax, [(x, y_from), (x, y_to)], color=color, lw=lw, style=style)


def _link_h(
    ax,
    x_from: float,
    x_to: float,
    y: float,
    *,
    color: str = "#2C2C2C",
    lw: float = 2.2,
    style: str = "-",
    arrow_end: bool = True,
) -> None:
    _draw_polyline(ax, [(x_from, y), (x_to, y)], color=color, lw=lw, style=style, arrow_end=arrow_end)


def draw_overview(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 10.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 11.2)
    ax.axis("off")
    ax.set_title(
        "jetson-demo 总体架构（Quantization-Activated Backdoor Demo）",
        fontsize=14,
        weight="bold",
        pad=12,
    )

    cx = 7.0

    # y positions top -> bottom (leave margin at bottom)
    y_ui = 9.6
    y_svc = 8.0
    y_hub = 6.2
    y_pipe = 4.2
    y_model = 2.0
    y_def = 0.35

    h_ui, h_svc, h_hub, h_pipe, h_model, h_def = 1.0, 1.15, 1.45, 1.55, 1.55, 0.85

    _box(
        ax,
        (0.6, y_ui),
        12.8,
        h_ui,
        "展示层  Browser Dashboard\n"
        "web/jetson_dashboard（默认）  |  web/react_dashboard（/react）  |  MJPEG / REST / WebSocket",
        "#E8F4FD",
        fontsize=9,
    )

    _box(
        ax,
        (0.6, y_svc),
        12.8,
        h_svc,
        "服务入口（二选一，共用 FrameHub / Pipeline）\n"
        "camera_web_preview.py（stdlib HTTP，现场默认）  |  camera_web_fastapi.py（FastAPI + /docs + /ws/status）",
        "#FFF4E5",
        fontsize=8.5,
    )

    _box(
        ax,
        (0.6, y_hub),
        12.8,
        h_hub,
        "运行时核心  FrameHub\n"
        "· 视频线程：OpenCV 采集（USB / CSI GStreamer / 图片·视频）→ MJPEG 编码\n"
        "· 推理线程（可选异步）：按 infer_every_n 处理最新帧 → 更新 status metrics\n"
        "· 控制面：mode(normal|triggered|defended)、attack、defense、defense_mode",
        "#E8F8E8",
        fontsize=9,
    )

    _box(
        ax,
        (0.6, y_pipe),
        12.8,
        h_pipe,
        "RealtimeQuraPipeline  →  demos/demo_qura_realtime_full.py\n"
        "Trigger 注入 → 选择 Backbone → ViT 推理 + Attention → 可选防御（oracle / regionblur / patchdrop）",
        "#F3E8FF",
        fontsize=9,
    )

    _box(ax, (0.6, y_model), 2.8, h_model, "FP32 ViT-B/16\n或 JIT Bundle", "#FDE8E8", fontsize=8)
    _box(ax, (3.55, y_model), 2.9, h_model, "INT8-QURA\n(MQBench W4A8)", "#FDE8E8", fontsize=8)
    _box(ax, (6.6, y_model), 2.8, h_model, "TRT logits\n(仅分类试点)", "#FDE8E8", fontsize=8)
    _box(ax, (9.55, y_model), 3.85, h_model, "FireViT-FP32\n(--fire-checkpoint)", "#E8F8E8", fontsize=8)

    _box(
        ax,
        (0.6, y_def),
        12.8,
        h_def,
        "防御库  defenses/regiondrop/region_detector.py\n"
        "AttentionHook · multi_scale_region_search · PatchDrop zero-mask",
        "#EEEEEE",
        fontsize=8.5,
    )

    # Main spine: box bottom -> next box top (solid, full length)
    bw = 12.8
    layers = [
        (0.6, y_ui, bw, h_ui),
        (0.6, y_svc, bw, h_svc),
        (0.6, y_hub, bw, h_hub),
        (0.6, y_pipe, bw, h_pipe),
    ]
    for i in range(len(layers) - 1):
        _link_v(ax, cx, _edge(layers[i], "bottom")[1], _edge(layers[i + 1], "top")[1], lw=2.4)

    # Pipeline -> backbone bus (complete T-junction)
    model_boxes = [
        (0.6, y_model, 2.8, h_model),
        (3.55, y_model, 2.9, h_model),
        (6.6, y_model, 2.8, h_model),
        (9.55, y_model, 3.85, h_model),
    ]
    pipe_box = (0.6, y_pipe, bw, h_pipe)
    pipe_bottom = _edge(pipe_box, "bottom")[1]
    model_top = _edge(model_boxes[0], "top")[1]
    bus_y = (pipe_bottom + model_top) / 2
    _link_v(ax, cx, pipe_bottom, bus_y, color="#5A5A5A", lw=2.0, style="--")
    bus_x0 = _edge(model_boxes[0], "bottom")[0]
    bus_x1 = _edge(model_boxes[3], "bottom")[0]
    _draw_polyline(
        ax,
        [(bus_x0, bus_y), (bus_x1, bus_y)],
        color="#5A5A5A",
        lw=2.0,
        style="--",
        arrow_end=False,
    )
    for mb in model_boxes:
        mx = _edge(mb, "bottom")[0]
        _link_v(ax, mx, bus_y, _edge(mb, "top")[1], color="#5A5A5A", lw=2.0, style="--")
    ax.text(cx, bus_y + 0.18, "调用 Backbone（FP32 / INT8-QURA / TRT / FireViT）", ha="center", fontsize=8.5, color="#444444")

    # INT8 backbone -> defense library (orthogonal, label on horizontal segment)
    mid_model = model_boxes[1]
    def_box = (0.6, y_def, bw, h_def)
    mid_x = _edge(mid_model, "bottom")[0]
    def_top_y = _edge(def_box, "top")[1]
    lane_y = (_edge(mid_model, "bottom")[1] + def_top_y) / 2
    _draw_polyline(
        ax,
        [
            (mid_x, _edge(mid_model, "bottom")[1]),
            (mid_x, lane_y),
            (cx, lane_y),
            (cx, def_top_y),
        ],
        color="#5A5A5A",
        lw=2.0,
        style="--",
    )
    ax.text((mid_x + cx) / 2, lane_y + 0.12, "引用防御库", ha="center", fontsize=8.5, color="#444444")

    fig.savefig(path, dpi=180, bbox_inches="tight", pad_inches=0.35, facecolor="white")
    plt.close(fig)


def draw_flow(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8.5)
    ax.axis("off")
    ax.set_title("实时推理数据流（Web Demo 主路径）", fontsize=14, weight="bold", pad=12)

    y_main = 4.8
    h = 1.25
    steps = [
        (0.3, y_main, 2.0, h, "视频源\nCSI/USB/文件", "#E8F4FD"),
        (2.5, y_main, 2.0, h, "FrameHub\n采集+MJPEG", "#E8F8E8"),
        (4.7, y_main, 2.0, h, "Trigger\n注入", "#FFF4E5"),
        (6.9, y_main, 2.1, h, "ViT/QURA\nForward", "#FDE8E8"),
        (9.2, y_main, 2.0, h, "Attention\n提取+检测", "#F3E8FF"),
        (11.4, y_main, 2.2, h, "Status/API\nprediction·top-k", "#E8F4FD"),
    ]
    for x, y, w, hb, t, c in steps:
        _box(ax, (x, y), w, hb, t, c, fontsize=8.5)

    boxes = [(x, y, w, hb) for x, y, w, hb, _, _ in steps]
    for i in range(len(boxes) - 1):
        p0 = _edge(boxes[i], "right")
        p1 = _edge(boxes[i + 1], "left")
        _link_h(ax, p0[0], p1[0], p0[1], lw=2.4)

    # Defense branch: under ViT, then orthogonal to Status (no diagonal)
    y_def = 2.3
    h_def = 1.35
    vit_box = boxes[3]
    status_box = boxes[5]
    def_box_geom = (5.4, y_def, 5.0, h_def)
    _box(
        ax,
        def_box_geom[:2],
        def_box_geom[2],
        def_box_geom[3],
        "防御分支（defense_on 时）\noracle / regionblur / patchdrop\n二次 forward → 恢复预测",
        "#E8F8E8",
        fontsize=8.5,
    )
    _box(
        ax,
        (0.3, y_def),
        4.6,
        h_def,
        "演示语义对照\nFP32+trigger → 后门休眠\nINT8-QURA+trigger → 后门激活\nDefense → 脱离目标类",
        "#FFF9E6",
        fontsize=8.5,
    )

    vit_bottom = _edge(vit_box, "bottom")
    def_top = _edge(def_box_geom, "top")
    _link_v(ax, vit_bottom[0], vit_bottom[1], def_top[1], lw=2.4)
    ax.text(vit_bottom[0] + 0.22, (vit_bottom[1] + def_top[1]) / 2, "可疑/激活", fontsize=8, color="#444444", va="center")

    def_right = _edge(def_box_geom, "right")
    status_bottom = _edge(status_box, "bottom")
    _draw_polyline(
        ax,
        [
            def_right,
            (status_bottom[0], def_right[1]),
            status_bottom,
        ],
        color="#1B7A3D",
        lw=2.4,
    )
    ax.text(
        (def_right[0] + status_bottom[0]) / 2,
        def_right[1] + 0.14,
        "更新预测",
        fontsize=8,
        color="#1B7A3D",
        ha="center",
    )

    ax.text(
        0.35,
        7.2,
        "离线保底：jetson_demo_imagenet.py（FP32 JIT 现场 + x86 预计算表回放）",
        fontsize=9,
        va="top",
        bbox=dict(boxstyle="round", facecolor="#F5F5F5", edgecolor="#999999"),
    )

    fig.savefig(path, dpi=180, bbox_inches="tight", pad_inches=0.35, facecolor="white")
    plt.close(fig)


def draw_offline(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 4.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    ax.set_title("基础版离线 Demo 架构（分离部署策略）", fontsize=13, weight="bold", pad=10)

    _box(
        ax,
        (0.5, 2.2),
        3.5,
        1.5,
        "x86 侧\nexport_jit_imagenet_vit.py\n预计算 INT8/防御表",
        "#E8F4FD",
        fontsize=9,
    )
    _box(ax, (4.5, 2.2), 3.0, 1.5, "产物\noutputs/jetson_imagenet_demo/\n.jit.pt + JSON 表", "#FFF4E5", fontsize=9)
    _box(
        ax,
        (8.0, 2.2),
        3.5,
        1.5,
        "Jetson 侧\njetson_demo_imagenet.py\nFP32 JIT 现场 + 表回放",
        "#E8F8E8",
        fontsize=9,
    )
    _link_h(ax, 4.0, 4.5, 2.95, lw=2.4)
    _link_h(ax, 7.5, 8.0, 2.95, lw=2.4)

    ax.text(
        0.5,
        0.5,
        "原因：MQBench 动态图与 global 状态导致 INT8-QURA 无法直接 TorchScript/JIT 导出",
        fontsize=9,
        color="#333333",
    )

    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _add_title(doc: Document, text: str, size: int = 16) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def append_to_docx(
    overview: Path,
    flow: Path,
    offline: Path,
) -> None:
    if BACKUP.exists():
        pass
    else:
        shutil.copy2(DOCX_SRC, BACKUP)

    doc = Document(str(DOCX_SRC))

    doc.add_page_break()
    _add_title(doc, "附录：Demo 框架架构设计图")

    p = doc.add_paragraph()
    p.add_run(f"图示生成日期：{date.today().isoformat()}。").italic = True
    doc.add_paragraph(
        "下图基于当前 jetson-demo 仓库（README、camera_web_preview.py、"
        "demo_qura_realtime_full.py、defenses/regiondrop）整理，"
        "用于说明「量化激活后门 + 推理时防御」演示框架的分层结构与数据流。"
    )

    _add_subtitle(doc, "A. 总体分层架构")
    doc.add_picture(str(overview), width=Cm(16.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bullets_a = [
        "展示层：浏览器通过 MJPEG（/stream.mjpg）看视频，通过 REST 或 WebSocket 读推理状态并下发控制。",
        "服务层：默认 stdlib HTTP（camera_web_preview.py）；可选 FastAPI 入口提供 /docs 与 /ws/status，推理语义不变。",
        "运行时：FrameHub 将「高帧率视频」与「低频 ViT 推理」解耦，保证现场约 20 FPS 级别的流畅预览。",
        "模型层（QURA 主线）：FP32/JIT 展示后门休眠；INT8-QURA（MQBench）展示量化后 trigger 激活；TRT 仅作 logits 延迟试点。",
        "模型层（火焰支线）：FireViT-FP32（--fire-checkpoint 模式），fire/no_fire 二分类，滑动窗口告警；三轮微调后视频测试 F1=97.44%，仅 1 FP。",
        "防御层：基于 CLS→patch attention 定位可疑区域，支持 oracle、regionblur、patchdrop 三种推理时缓解（QURA 主线）。",
    ]
    for item in bullets_a:
        doc.add_paragraph("• " + item)

    _add_subtitle(doc, "B. 实时 Web Demo 数据流")
    doc.add_picture(str(flow), width=Cm(16.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bullets_b = [
        "控制语义：Normal 优先 FP32；Triggered 开启 trigger 并走 INT8-QURA；Defended 同时开启防御。",
        "Trigger 在 normalized tensor 空间注入，与离线 ImageNet 流程一致；页面红框映射到真实模型输入位置。",
        "PatchDrop 在 INT8 路径上对可疑 patch 做 zero-mask 后二次 forward；oracle/regionblur 在 tensor 上恢复/模糊触发区域。",
        "若使用 --int8-only，Normal 模式可能 fallback 到 INT8，演示 FP32 dormant 时需加载 FP32 或 JIT bundle。",
    ]
    for item in bullets_b:
        doc.add_paragraph("• " + item)

    _add_subtitle(doc, "C. 离线保底 Demo 架构")
    doc.add_picture(str(offline), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "离线路线适合汇报保底与无摄像头环境：Jetson 只跑 FP32 JIT，"
        "INT8 激活与防御效果由 x86 预计算结果展示，避免在设备端完整加载 MQBench。"
    )

    _add_subtitle(doc, "D. 目录与模块对照")
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "路径"
    hdr[1].text = "职责"
    rows = [
        ("demos/demo_qura_realtime_full.py", "实时 ViT/QURA 推理、trigger、attention、防御主逻辑"),
        ("scripts/camera_web_preview.py", "FrameHub + HTTP/MJPEG + 默认 dashboard；--fire-checkpoint 启用 FireBackbone 火焰检测"),
        ("scripts/camera_web_fastapi.py", "FastAPI 包装，WebSocket 状态推送"),
        ("defenses/regiondrop/region_detector.py", "AttentionHook、区域搜索、PatchDrop"),
        ("third_party/qura/", "QuRA / MQBench 量化与 checkpoint"),
        ("deploy/trt_runner.py", "可选 TensorRT logits 推理"),
        ("web/jetson_dashboard/", "默认静态控制台"),
        ("web/react_dashboard/", "React ES Module 预览控制台"),
        ("scripts/finetune_lab_fire_vit.py", "火焰 ViT 微调脚本（冻结 backbone + 线性头）"),
        ("scripts/test_vit_on_videos.py", "视频滑动窗口评估，生成 summary_v2.json"),
        ("outputs/lab_fire_vit/", "微调权重 lab_fire_vit_head_best.pt + metrics"),
    ]
    for path, role in rows:
        row = table.add_row().cells
        row[0].text = path
        row[1].text = role

    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run(
        "说明：架构图描述的是当前仓库已实现的主线能力。"
        "RT-DETR、BadDet+、CleAnSight 等方向在文档第 10 节中记为后续扩展，未纳入本图。"
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.save(str(DOCX_DST))


def replace_docx_images(overview: Path, flow: Path, offline: Path) -> None:
    """Replace appendix PNGs inside existing docx (image1/2/3)."""
    import tempfile
    import zipfile

    mapping = {
        "word/media/image1.png": overview,
        "word/media/image2.png": flow,
        "word/media/image3.png": offline,
    }
    tmp = Path(tempfile.mkdtemp()) / "docx_new.zip"
    with zipfile.ZipFile(DOCX_DST, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in mapping:
                data = mapping[item.filename].read_bytes()
            zout.writestr(item, data)
    shutil.copy2(DOCX_DST, BACKUP)
    shutil.move(str(tmp), str(DOCX_DST))


def main() -> None:
    import sys

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    overview = OUT_DIR / "arch_overview.png"
    flow = OUT_DIR / "arch_flow.png"
    offline = OUT_DIR / "arch_offline.png"

    draw_overview(overview)
    draw_flow(flow)
    draw_offline(offline)

    def _print(msg: str) -> None:
        sys.stdout.buffer.write((msg + "\n").encode("utf-8"))
        sys.stdout.flush()

    if "--replace-images" in sys.argv:
        replace_docx_images(overview, flow, offline)
        _print(f"Replaced images in: {DOCX_DST}")
    elif "--append" in sys.argv:
        append_to_docx(overview, flow, offline)
        _print(f"Appended appendix to: {DOCX_DST}")
    else:
        replace_docx_images(overview, flow, offline)
        _print(f"Replaced images in: {DOCX_DST} (use --append to add appendix again)")

    _print(f"Saved diagrams to {OUT_DIR}")
    _print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
